# Проверка титула.
from docx import Document


# import picture
# import section
# import styleFont
from src.typeCheck import styleFont, picture


document = Document('src/example/font1.docx')

# def document_verification(document):
#     for paragraph in document.paragraphs:
#         # styleFont.font(paragraph)
#         pass

picture.extract_images_from_docx(document)



# document_verification(document)
# # (Поля) Проверка полей страниц.
# section.indent(document)
#
# # (Текст) Проверка текста
# styleFont.font(document) # [Внести списки]
#
# # (Текст) Проверка заголовков
# styleFont.header(document)
#
# # (Страницы) Структура
# section.structure(document) # [Доделать нумерацию]
#
# # (Илюстрация) Проверка изображений
# picture.extract_images_from_docx(document) # [Доделать нумерацию]
#
# # (Страницы) Нумерация
# section.numbering(document)

# (Текст) Перечисления
# styleFont.find_lists(document)

# picture.explore_xml_structure(document)




















# import picture
# from style import Style
# from errorList import Error
# # Проверка шапки титула (Образовательное учреждение).
# list_par_head = [document.paragraphs[0], document.paragraphs[1], document.paragraphs[2]]
# list_par_head_cor = ['Министерство науки и высшего образования Российской Федерации',
#                      'ФГАОУ ВО «УрФУ имени первого Президента России Б.Н. Ельцина»',
#                      'Кафедра «школа бакалавриата (школа)»']
#
# for i in range(len(list_par_head)):
#     par_h = list_par_head[i]
#     cor_h = list_par_head_cor[i]
#     if i != 2:
#         if par_h.text != cor_h or not Style.isBold(par_h) or Style.isItalic(par_h) or Style.alignment(par_h, 1):
#             print(Error.check(par_h.text, 10))
#
#     else:
#         if par_h.text != cor_h or Style.isBold(par_h) or Style.isItalic(par_h) or Style.alignment(par_h, 1):
#             print(Error.check(par_h.text, 10))
#
#
# # Проверка Оценки работы.
# list_par_score = [document.paragraphs[8], document.paragraphs[9]]
# list_par_score_cor = ["Оценка работы______________", "Руководитель от УрФУ"]
#
# for i in range(len(list_par_score)):
#     par_s = list_par_score[i]
#     cor_s = list_par_score_cor[i]
#     if i == 0:
#         if cor_s != par_s.text or Style.isBold(par_s) or Style.isItalic(par_s) or Style.alignment(par_s, 2):
#             print(Error.check(par_s.text, 11))
#     else:
#         if cor_s != par_s.text[:len(cor_s)] or Style.isBold(par_s) or Style.isItalic(par_s) or Style.alignment(par_s, 2):
#             print(Error.check(par_s.text, 11))
















