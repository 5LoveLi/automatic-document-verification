from docx import Document

from src.typeCheck import styleFont

# документ - 'src/example/font1.docx' содержит все представленные ошибки 
# когда кегель равен 11 программа не добавляет ошибку кегеля, с лругими размерами работает 

def test_headset():
    error = 'Ошибка в оформление текста: Некоректный шрифт. Необходимо использовать - Times New Roman.'
    assert error in styleFont.font(Document('src/example/font1.docx')) 

def test_kegel():
    error = 'Ошибка в оформление текста: Некоректный кегель. Необходимо использовать - 14.'
    assert error in styleFont.font(Document('src/example/font1.docx'))

def test_color():
    error = 'Ошибка в оформление текста: Некорректный цвет. Необходимо использовать - черный.'
    assert error in styleFont.font(Document('src/example/font1.docx'))

def test_outline():
    error = 'Ошибка в оформление текста: Некорректное начертание. Необходимо использовать - прямое начертание.'
    assert error in styleFont.font(Document('src/example/font1.docx'))

def test_interval():
    error = 'Ошибка в оформление текста: Некорректный интервал. Необходимо использовать - 1,5.'
    assert error in styleFont.font(Document('src/example/font1.docx'))