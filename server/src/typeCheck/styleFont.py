import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from src import errorList



def font(doc):
    error_list = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name != 'Heading 1' and paragraph.style.name != 'Heading 2' and paragraph.style.name != 'Heading 3':

            for run in paragraph.runs:
                font_name = run.font.name

                if font_name is not None and font_name != 'Times New Roman':
                    error_list.append(errorList.error[13])
                    print("\033[31m {}\033[0m" .format(errorList.error[13]))

                font_size = run.font.size
                if font_size is not None and font_size.pt != 14:
                    error_list.append(errorList.error[14])
                    print("\033[31m {}\033[0m" .format(errorList.error[14]))

                font_color = run.font.color.rgb
                if font_color is not None and font_color != (0, 0, 0):
                    error_list.append(errorList.error[16])
                    print("\033[31m {}\033[0m" .format(errorList.error[16]))

                if not(not run.font.italic and not run.font.bold and not run.font.underline):
                    error_list.append(errorList.error[17])
                    print("\033[31m {}\033[0m" .format(errorList.error[17]))

            if paragraph.paragraph_format.line_spacing != 1.5:
                p = paragraph.paragraph_format.line_spacing
                error_list.append(errorList.error[15])
                print("\033[31m {}\033[0m" .format(errorList.error[15]))
    
    return error_list


def header(doc):
    # Проверка заголовков разделов
    # section_num = 1
    error_list = []
    for paragraph in doc.paragraphs:
        alignment = paragraph.alignment
        if paragraph.style.name == 'Heading 1' and alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            section_text = paragraph.text.strip()

            # Проверка отсутствия точки в конце
            if section_text.endswith('.'):
                error_list.append(errorList.error[18])
                print("\033[31m {}\033[0m" .format(errorList.error[18]))

            # Проверка строчных букв
            if not any(letter.islower() for letter in section_text):
                error_list.append(errorList.error[19])
                print("\033[31m {}\033[0m" .format(errorList.error[19]))

            if section_text[0].islower():
                error_list.append(errorList.error[20])
                print("\033[31m {}\033[0m" .format(errorList.error[20]))

            # Проверка полужирного шрифта
            if not paragraph.runs[0].bold:
                error_list.append(errorList.error[21])
                print("\033[31m {}\033[0m" .format(errorList.error[21]))

            # Проверка расположения слева
            if paragraph.alignment != None:
                error_list.append(errorList.error[22])
                print("\033[31m {}\033[0m" .format(errorList.error[22]))

            # Проверка отсутствия подчеркивания
            if paragraph.runs[0].underline:
                error_list.append(errorList.error[23])
                print("\033[31m {}\033[0m" .format(errorList.error[23]))

            # Проверка цвета заголовка
            if paragraph.runs[0].font.color.rgb != RGBColor(0, 0, 0):
                error_list.append(errorList.error[16])
                print("\033[31m {}\033[0m" .format(errorList.error[16]))

            # Проверка абзацного отступа
            if paragraph.paragraph_format.first_line_indent != 450215:
                error_list.append(errorList.error[24])
                print("\033[31m {}\033[0m" .format(errorList.error[24]))
                # print("заголовков разделов")

            if paragraph.paragraph_format.line_spacing != 1.5:
                p = paragraph.paragraph_format.line_spacing
                error_list.append(errorList.error[15])
                print("\033[31m {}\033[0m" .format(errorList.error[15]))

            for run in paragraph.runs:
                font_name = run.font.name
                if font_name != 'Times New Roman':
                    error_list.append(errorList.error[13])
                    print("\033[31m {}\033[0m" .format(errorList.error[13]))

            # Проверка нумерации
            # if section_text[0] == str(section_num) + " ":
            #     print('Проверка нумерации')
            # print(section_text)
            # section_num += 1

            # print('Проверен 1-ый заголовок')
        continue

    return error_list


    # # Проверка заголовков подразделов
    # for paragraph in doc.paragraphs:
    #     if paragraph.style.name == 'Heading 2':
    #         section_text = paragraph.text.strip()

    #         # Проверка отсутствия точки в конце
    #         if section_text.endswith('.'):
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',18)))

    #         # Проверка строчных букв
    #         if not any(letter.islower() for letter in section_text):
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',19)))

    #         if section_text[0].islower():
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',20)))

    #         # Проверка полужирного шрифта
    #         if not paragraph.runs[0].bold:
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',21)))

    #         # Проверка расположения слева
    #         if paragraph.alignment != None:
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('справа','#',22)))

    #         # Проверка отсутствия подчеркивания
    #         if paragraph.runs[0].underline:
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('справа','#',23)))

    #         # Проверка цвета заголовка
    #         if paragraph.runs[0].font.color.rgb != RGBColor(0, 0, 0):
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',16)))

    #         # Проверка абзацного отступа
    #         if paragraph.paragraph_format.first_line_indent != 450215:
    #             print("\033[31m {}\033[0m" .format(errorList.Error.check('#','#',24)))
    #             # print("заголовков подразделов")

    #         if paragraph.paragraph_format.line_spacing != 1.5:
    #             p = paragraph.paragraph_format.line_spacing
    #             print("\033[31m {}\033[0m" .format((errorList.Error.check(p, '#', 15))))

    #         for run in paragraph.runs:
    #             font_name = run.font.name
    #             if font_name != 'Times New Roman':
    #                 print("\033[31m {}\033[0m" .format(errorList.Error.check(font_name,'#',13)))

    #         # print('Проверен 2-ой заголовок')
    #     continue

    # # Проверка основного заголовка
    # for paragraph in doc.paragraphs:
    #     alignment = paragraph.alignment
    #     if paragraph.style.name == 'Heading 1' and alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
    #         section_text = paragraph.text.strip()

    #         # Проверка отсутствия точки в конце
    #         if section_text.endswith('.'):
    #             print(errorList.Error.check('#','#',18))

    #         # Проверка строчных букв
    #         if not section_text.isupper():
    #             print(errorList.Error.check('#','#',25))

    #         # Проверка полужирного шрифта
    #         if not paragraph.runs[0].bold:
    #             print(errorList.Error.check('#','#',21))

    #         # Проверка отсутствия подчеркивания
    #         if paragraph.runs[0].underline:
    #             print(errorList.Error.check('справа','#',23))

    #         # Проверка цвета заголовка
    #         if paragraph.runs[0].font.color.rgb != RGBColor(0, 0, 0):
    #             print(errorList.Error.check('#','#',16))

    #         for run in paragraph.runs:
    #             font_name = run.font.name
    #             if font_name != 'Times New Roman':
    #                 print(errorList.Error.check(font_name,'#',13))

    #         # print('Проверен 1-ый основной заголовок')
    #     continue



def find_lists(doc):
    # print(errorList.error[13])
    for paragraph in doc.paragraphs:
        if (
            paragraph.style.name.startswith('List')
            and paragraph.style.type == 1
            and paragraph.style.paragraph_format.left_indent is not None
            and paragraph.style.paragraph_format.left_indent > 0
        ):
            print(paragraph.text)
